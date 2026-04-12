"""Script to generate test data files for parser tests."""

from pathlib import Path
from docx import Document

# Create test data directory
data_dir = Path(__file__).parent / "data"
data_dir.mkdir(exist_ok=True)


# Create a simple valid DOCX file
def create_simple_docx():
    doc = Document()
    doc.add_heading("Test Resume", 0)
    doc.add_paragraph("John Doe")
    doc.add_paragraph("Email: john.doe@example.com")
    doc.add_paragraph("Phone: +1-234-567-8900")

    doc.add_heading("Experience", level=1)
    doc.add_paragraph("Software Engineer at Tech Company")
    doc.add_paragraph("Worked on various projects using Python, Java, and JavaScript.")

    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Python, Java, JavaScript, SQL, Git")

    # Save the document
    doc.save(str(data_dir / "valid_resume.docx"))
    print(f"Created valid_resume.docx")


# Create a DOCX with tables
def create_docx_with_tables():
    doc = Document()
    doc.add_heading("Resume with Tables", 0)

    # Add a table
    table = doc.add_table(rows=3, cols=2)
    table.style = "Light Grid Accent 1"

    # Fill the table
    cells = [
        ("Name", "Jane Smith"),
        ("Email", "jane@example.com"),
        ("Phone", "+1-987-654-3210"),
    ]

    for i, (label, value) in enumerate(cells):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    doc.add_paragraph()
    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Python, Data Analysis, Machine Learning")

    doc.save(str(data_dir / "with_tables.docx"))
    print(f"Created with_tables.docx")


# Create a DOCX with only whitespace/empty paragraphs
def create_empty_content_docx():
    doc = Document()
    doc.add_paragraph("")
    doc.add_paragraph("   ")
    doc.add_paragraph("\n")
    doc.save(str(data_dir / "empty_content.docx"))
    print(f"Created empty_content.docx")


if __name__ == "__main__":
    try:
        create_simple_docx()
        create_docx_with_tables()
        create_empty_content_docx()
        print("\nAll test DOCX files created successfully!")
    except Exception as e:
        print(f"Error creating test files: {e}")
