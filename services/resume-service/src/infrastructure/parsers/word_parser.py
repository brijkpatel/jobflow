from __future__ import annotations

from io import BytesIO

from docx import Document

from domain.exceptions import FileParsingError


class WordParser:
    def parse(self, file_bytes: bytes, filename: str) -> str:
        try:
            document = Document(BytesIO(file_bytes))
        except Exception as exc:  # pragma: no cover - python-docx error surface
            raise FileParsingError(f"failed to parse {filename}") from exc
        lines: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                lines.append(text)
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))
        if not lines:
            raise FileParsingError(f"no text extracted from {filename}")
        return "\n".join(lines)
